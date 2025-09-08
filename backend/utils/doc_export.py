#  Standard Library
import io
import re
from typing import Dict

# Third-Party Libraries
import openpyxl
from docx import Document
from docx.shared import Pt, RGBColor
from markdown_it import MarkdownIt
from mdit_py_plugins.front_matter import front_matter_plugin
from reportlab.lib import colors
from reportlab.lib.enums import TA_JUSTIFY
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle

# Internal Modules
from backend.utils.markdown import convert_markdown_bold



def add_markdown_to_doc(doc: Document, text: str):
    """
    Adds content to a .docx document, correctly handling Markdown syntax.

    This function parses the Markdown text and adds elements to the document,
    including paragraphs with bold text and tables. It now includes
    error handling to log issues during parsing.

    Args:
        doc: The python-docx Document object.
        text: The text, which may contain Markdown syntax.
    """
    # Configure the Markdown parser to handle tables.
    md = MarkdownIt().use(front_matter_plugin).enable('table')
    try:
        tokens = md.parse(text)
    except Exception as e:
        print(f"Error parsing markdown: {e}")
        doc.add_paragraph(f"Error parsing content: {text}")
        return

    i = 0
    while i < len(tokens):
        token = tokens[i]

        try:
            # Handle tables
            if token.type == "table_open":
                headers = []
                rows = []
                i += 1
                
                # Extract header and row data
                while i < len(tokens) and tokens[i].type != "table_close":
                    if tokens[i].type == "tr_open":
                        row_data = []
                        i += 1
                        while i < len(tokens) and tokens[i].type != "tr_close":
                            # Process cells (th and td)
                            if tokens[i].type in ("th_open", "td_open"):
                                i += 1
                                # Look for the inline token that contains the cell content
                                if i < len(tokens) and tokens[i].type == "inline":
                                    cell_text = tokens[i].content
                                    row_data.append(cell_text.strip())
                                i += 1
                                # The inline token is followed by the closing tag, so we need to move past both
                            i += 1
                        # The first row is always the header
                        if not headers:
                            headers = row_data
                        else:
                            rows.append(row_data)
                    i += 1

                # Build Word table from extracted data
                if headers:
                    table = doc.add_table(rows=len(rows) + 1, cols=len(headers))
                    table.style = "Table Grid"
                    
                    # Add headers
                    for j, header in enumerate(headers):
                        table.rows[0].cells[j].text = header
                    
                    # Add rows
                    for r, row_data in enumerate(rows):
                        for c, cell_text in enumerate(row_data):
                            table.rows[r + 1].cells[c].text = cell_text

            # Handle paragraphs
            elif token.type == "paragraph_open":
                paragraph = doc.add_paragraph()
                i += 1
                while i < len(tokens) and tokens[i].type != "paragraph_close":
                    if tokens[i].type == "inline":
                        # The content is within the inline token
                        for sub_token in tokens[i].children:
                            if sub_token.type == "text":
                                paragraph.add_run(sub_token.content)
                            elif sub_token.type == "strong_open":
                                # Handle bold text
                                strong_text = ""
                                strong_idx = tokens[i].children.index(sub_token) + 1
                                while strong_idx < len(tokens[i].children) and tokens[i].children[strong_idx].type == "text":
                                    strong_text += tokens[i].children[strong_idx].content
                                    strong_idx += 1
                                run = paragraph.add_run(strong_text)
                                run.bold = True
                    i += 1

                paragraph.paragraph_format.space_after = Pt(11)
                paragraph.paragraph_format.line_spacing = 1
            
            # The outer loop must always increment to avoid an infinite loop
            # and move to the next token, regardless of type
            i += 1
            
        except Exception as e:
            print(f"Error processing token type '{token.type}' at index {i}: {e}")
            doc.add_paragraph(f"Error processing content near: {token.content}")
            i += 1



def create_word_from_sections(form_data: Dict, ordered_sections: Dict) -> Document:
    """
    Generates a .docx document from proposal data.

    Args:
        form_data: A dictionary containing the proposal's metadata.
        ordered_sections: A dictionary of the proposal sections and their content.

    Returns:
        A python-docx Document object.
    """
    doc = Document()

    # --- Start of new formatting code ---
    # Set the style for the main title (Heading 1)
    style1 = doc.styles['Heading 1']
    style1.font.name = 'Arial'
    style1.font.size = Pt(16)
    style1.font.color.rgb = RGBColor(0x00, 0x72, 0xbc)

    # Set the style for section headings (Heading 2)
    style2 = doc.styles['Heading 2']
    style2.font.name = 'Arial'
    style2.font.size = Pt(14)
    style2.font.color.rgb = RGBColor(0x00, 0x72, 0xbc)

    # Set the style for the normal body text
    normal_style = doc.styles['Normal']
    normal_style.font.name = 'Arial'
    normal_style.font.size = Pt(11)
    # --- End of new formatting code ---

    # Add the main title.
    doc.add_heading("Draft Project Proposal", level=1)

    # Add the form data as a table.
    doc.add_heading("Project Prompt Details", level=2)
    try:
        table = doc.add_table(rows=1, cols=2)
        table.style = "Table Grid"
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = "Field"
        hdr_cells[1].text = "Value"
        for key, value in form_data.items():
            row_cells = table.add_row().cells
            row_cells[0].text = key
            row_cells[1].text = str(value)
    except Exception as e:
        print(f"Error adding form data table: {e}")

    # Add each proposal section to the document.
    for section, content in ordered_sections.items():
        if not content:
            continue  # Skip empty sections.
        doc.add_heading(section, level=2)
        add_markdown_to_doc(doc, content)

    return doc


def create_excel_from_sections(ordered_sections: Dict) -> bytes:
    """
    Generates an .xlsx document from proposal data, with each table in a separate sheet.

    Args:
        ordered_sections: A dictionary of the proposal sections and their content.

    Returns:
        A byte stream containing the generated .xlsx file.
    """
    workbook = openpyxl.Workbook()
    workbook.remove(workbook.active)  # Remove the default sheet

    md = MarkdownIt().use(front_matter_plugin).enable('table')
    table_count = 0

    for section, content in ordered_sections.items():
        if not content:
            continue

        tokens = md.parse(content)
        i = 0
        while i < len(tokens):
            token = tokens[i]
            if token.type == "table_open":
                table_count += 1
                headers = []
                rows = []
                i += 1  # Move to thead_open

                while i < len(tokens) and tokens[i].type != "table_close":
                    if tokens[i].type == "tr_open":
                        row_data = []
                        i += 1
                        while i < len(tokens) and tokens[i].type != "tr_close":
                            if tokens[i].type in ("th_open", "td_open"):
                                i += 1
                                if i < len(tokens) and tokens[i].type == "inline":
                                    cell_text = tokens[i].content
                                    row_data.append(cell_text.strip())
                                i += 1  # Move past closing tag
                            i += 1
                        if not headers:
                            headers = row_data
                        else:
                            rows.append(row_data)
                    i += 1

                if headers:
                    sheet_title = f"Table {table_count}"
                    # Truncate sheet title if it's too long
                    if len(sheet_title) > 31:
                        sheet_title = sheet_title[:31]
                    worksheet = workbook.create_sheet(title=sheet_title)
                    worksheet.append(headers)
                    for row in rows:
                        worksheet.append(row)
            i += 1

    excel_buffer = io.BytesIO()
    workbook.save(excel_buffer)
    excel_buffer.seek(0)
    return excel_buffer.read()


def create_pdf_from_sections(form_data: Dict, ordered_sections: Dict) -> bytes:
    """
    Generates a PDF document from proposal data using ReportLab and returns it as a byte buffer.

    Args:
        form_data: A dictionary containing the proposal's metadata.
        ordered_sections: A dictionary of the proposal sections and their content.

    Returns:
        A byte buffer containing the generated PDF.
    """
    buffer = io.BytesIO()
    doc = SimpleDocTemplate(
        buffer, pagesize=A4, rightMargin=72, leftMargin=72, topMargin=72, bottomMargin=72
    )
    styles = getSampleStyleSheet()

    # Define a custom style for justified body text.
    normal_style = ParagraphStyle(name="Body", parent=styles["Normal"], alignment=TA_JUSTIFY, leading=14)

    story = []

    # Add the main title.
    story.append(Paragraph("Project Proposal", styles["Title"]))
    story.append(Spacer(1, 12))

    # Create a table for the form data.
    table_data = [["Field", "Value"]]
    for key, value in form_data.items():
        table_data.append([key, str(value)])

    table = Table(table_data, colWidths=[2.5 * inch, 3.5 * inch])
    table.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (-1, 0), colors.lightgrey),
                ("TEXTCOLOR", (0, 0), (-1, 0), colors.black),
                ("ALIGN", (0, 0), (-1, -1), "LEFT"),
                ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
                ("FONTSIZE", (0, 0), (-1, -1), 10),
                ("BOTTOMPADDING", (0, 0), (-1, 0), 8),
                ("GRID", (0, 0), (-1, -1), 0.5, colors.grey),
            ]
        )
    )
    story.append(table)
    story.append(Spacer(1, 20))

    # Add each proposal section to the document.
    for section, content in ordered_sections.items():
        if not content:
            continue  # Skip empty sections.

        story.append(Paragraph(section, styles["Heading2"]))
        story.append(Spacer(1, 6))

        # Split content into paragraphs and format them.
        for paragraph in content.split("\n\n"):
            # Convert markdown bold to reportlab's <b> tag.
            cleaned_paragraph = convert_markdown_bold(paragraph.strip())
            story.append(Paragraph(cleaned_paragraph, normal_style))
            story.append(Spacer(1, 10))

    # Build the PDF document.
    doc.build(story)
    pdf_bytes = buffer.getvalue()
    buffer.close()
    return pdf_bytes


def generate_final_markdown(generated_sections: Dict) -> str:
    """
    Compiles all generated sections into a single Markdown string.

    Args:
        generated_sections: A dictionary of the proposal sections and their content.

    Returns:
        A string containing the full proposal in Markdown format.
    """
    markdown_content = ""
    for section, content in generated_sections.items():
        markdown_content += f"## {section}\n\n{content.strip()}\n\n"
    return markdown_content
