#TO RUN THIS CODE - uvicorn main:app --host 172.1.24.95 --port 8501 --reload

#  Standard Library
import json
import os
import re
import uuid
import io
from datetime import datetime, timedelta
import pprint
import traceback

#  Third-Party Libraries
from fastapi import FastAPI, HTTPException, File, UploadFile, Form
from fastapi.responses import FileResponse, JSONResponse
from fastapi import Request
from fastapi.responses import Response
from fastapi import Query
from pydantic import BaseModel
from typing import Dict, Optional
import redis
from fastapi.middleware.cors import CORSMiddleware
from werkzeug.security import generate_password_hash, check_password_hash
import jwt
from functools import wraps
from sqlalchemy import create_engine, text
import urllib.parse
from fastapi import Depends

#  Internal Modules
from crew import ProposalCrew

#  Document Handling (python-docx)
from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

from fastapi.responses import FileResponse
from docx import Document
from docx.shared import Pt
# from docx2pdf import convert
import tempfile

from reportlab.lib.pagesizes import A4
from reportlab.pdfgen import canvas
from reportlab.lib.units import inch

from reportlab.platypus import (
    SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
)
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.enums import TA_JUSTIFY
from reportlab.lib import colors
from reportlab.lib.pagesizes import A4
from reportlab.lib.units import inch
import re
from docx.shared import Pt

from apscheduler.schedulers.background import BackgroundScheduler
from datetime import timedelta, datetime, timezone


# import platform
# import pypandoc
# from docx2pdf import convert

app = FastAPI()

# Allow CORS from specific frontend origin(s)
origins = [
    "https://proposal-drafter.azurewebsites.net/",
     "http://localhost:8503" # here the frontend url will be there
]

app.add_middleware(
    CORSMiddleware,
    allow_origins=origins,
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)



session_data = {}

redis_client = redis.Redis(host="redis", port=6379, db=0, decode_responses=True)

SECRET_KEY = os.getenv("SECRET_KEY", "your_default_dev_secret")

db_username = os.getenv('DB_USERNAME')
db_password = os.getenv('DB_PASSWORD')
db_host = os.getenv('DB_HOST')
db_port = os.getenv('DB_PORT')
db_name = os.getenv('DB_NAME')
                
encoded_password = urllib.parse.quote_plus(db_password)
    
connection_string = f"postgresql://{db_username}:{encoded_password}@{db_host}:{db_port}/{db_name}"
engine = create_engine(connection_string)

def get_current_user(request: Request):
    token = request.cookies.get("auth_token")
    if not token:
        raise HTTPException(status_code=401, detail="Authentication token missing.")

    try:
        payload = jwt.decode(token, SECRET_KEY, algorithms=["HS256"])
        email = payload.get("email")
        if not email:
            raise HTTPException(status_code=401, detail="Invalid token payload.")

        with engine.connect() as connection:
            result = connection.execute(text("SELECT id, name, email FROM users WHERE email = :email"), {"email": email})
            user = result.fetchone()

        if not user:
            raise HTTPException(status_code=404, detail="User not found.")

        return {
            "user_id": str(user[0]),
            "name" : user[1],
            "email": user[2]
        }

    except jwt.ExpiredSignatureError:
        raise HTTPException(status_code=401, detail="Token expired.")
    except jwt.InvalidTokenError:
        raise HTTPException(status_code=401, detail="Invalid token.")
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Auth error: {str(e)}")


# Initialize Redis with error handling
try:
    redis_client = redis.Redis(host="redis", port=6379, db=0, decode_responses=True)
    # Test connection
    redis_client.ping()
    print("Successfully connected to Redis")
except redis.ConnectionError:
    print("Warning: Could not connect to Redis. Using in-memory storage as fallback.")
    # Use a simple dict as fallback if Redis is not available
    class DictStorage:
        def __init__(self):
            self.storage = {}
            
        def setex(self, key, ttl, value):
            self.storage[key] = value
            
        def set(self, key, value):
            self.storage[key] = value
            
        def get(self, key):
            return self.storage.get(key)
        
        def delete(self, key):
            self.storage.pop(key, None)
    
    redis_client = DictStorage()



CONFIG_PATH = "config/templates/iom_proposal_template.json"
SECTIONS = ["Summary", "Rationale", "Project Description", "Partnerships and Coordination", "Monitoring", "Evaluation"]

# Load JSON configuration
with open(CONFIG_PATH, "r", encoding="utf-8") as file:
    proposal_data = json.load(file)

# Store processed sections
generated_sections = {}

class BaseDataRequest(BaseModel):
    form_data: Dict[str, str]  
    project_description: str 

class SectionRequest(BaseModel):
    section: str
    proposal_id: str  

class RegenerateRequest(BaseModel):
    section: str
    concise_input: str
    proposal_id: str

class SaveDraftRequest(BaseModel):
    session_id: Optional[str] = None 
    proposal_id: Optional[str] = None
    form_data: Dict[str, str]
    project_description: str
    generated_sections: Optional[Dict[str, str]] = {}

class FinalizeProposalRequest(BaseModel):
    proposal_id: str

def get_session_id():
    """Generate a unique session ID (UUID) for each user session."""
    return str(uuid.uuid4())  # Later, replace with SSO session ID

@app.post("/api/store_base_data")
async def store_base_data(request: BaseDataRequest,current_user: dict = Depends(get_current_user) ):
    session_id = get_session_id()  
    data = {
        "form_data": request.form_data,
        "project_description": request.project_description,
        "user_id": current_user["user_id"] 
    }
    # Store in Redis with a 1-hour expiration (3600 seconds)
    redis_client.setex(session_id, 3600, json.dumps(data))
    
    return {"message": "Base data stored successfully", "session_id": session_id}


@app.get("/api/get_base_data/{session_id}")
async def get_base_data(session_id: str):
    data = redis_client.get(session_id)
    if data is None:
        raise HTTPException(status_code=404, detail="Session data not found")
    
    return json.loads(data)  # Convert JSON string back to dictionary

def regenerate_section_logic(session_id: str, section: str, concise_input: str, proposal_id: str) -> str:
    """Shared logic for regenerating a section using a custom concise input (from user or evaluator)."""
    session_data_str = redis_client.get(session_id)

    if not session_data_str:
        raise HTTPException(status_code=400, detail="Base data not found. Please store it first.")

    session_data = json.loads(session_data_str)
    form_data = session_data["form_data"]
    project_description = session_data["project_description"]

    section_data = next((s for s in proposal_data["sections"] if s["section_name"] == section), {})
    if not section_data:
        raise HTTPException(status_code=400, detail=f"Invalid section name: {section}")

    instructions = section_data.get("instructions", "")
    word_limit = section_data.get("word_limit", 350)

    proposal_crew = ProposalCrew()
    crew_instance = proposal_crew.regenerate_proposal_crew()

    section_input = {
        "section": section,
        "form_data": form_data,
        "project_description": project_description,
        "instructions": instructions,
        "word_limit": word_limit,
        "concise_input": concise_input
    }

    result = crew_instance.kickoff(inputs=section_input)
    raw_output = result.raw.replace("`", "")
    raw_output = re.sub(r'[\x00-\x1F\x7F]', '', raw_output)  
    print("raw_output==>",raw_output)
    parsed = json.loads(raw_output)

    try:
        parsed = json.loads(raw_output)
        generated_text = parsed.get("generated_content", "").strip()
        evaluation_status = parsed.get("evaluation_status", "")
        feedback = parsed.get("feedback", "")
    except Exception as e:
        return JSONResponse(
            status_code=500,
            content={"error": "Invalid JSON response from crew", "details": str(e)}
        )

    if not result:
        raise HTTPException(status_code=500, detail=f"Failed to regenerate content for {section}")

    if "generated_sections" not in session_data:
        session_data["generated_sections"] = {}

    session_data["generated_sections"][section] = generated_text
    redis_client.set(session_id, json.dumps(session_data))
    
    user_id = session_data.get("user_id")
    if not user_id:
        print("[regenerate_section_logic] ❌ user_id missing in Redis. Skipping DB update.")
        return generated_text
    
    try:
        with engine.begin() as connection:
            result = connection.execute(
                text("SELECT generated_sections FROM proposals WHERE id = :proposal_id AND user_id = :user_id"),
                {
                    "proposal_id": proposal_id,
                    "user_id": session_data.get("user_id")
                }
            )
            draft = result.fetchone()

            if draft:
                updated_sections = json.loads(draft[0]) if draft[0] else {}
                updated_sections[section] = generated_text

                connection.execute(
                    text("""
                        UPDATE proposals
                        SET generated_sections = :sections, updated_at = NOW()
                        WHERE id = :proposal_id AND user_id = :user_id
                    """),
                    {
                        "sections": json.dumps(updated_sections),
                        "proposal_id": proposal_id,
                        "user_id": session_data["user_id"]
                    }
                )
    except Exception as e:
        print(f"[DB UPDATE ERROR in regenerate_section_logic] {e}")


    return generated_text


@app.post("/api/process_section/{session_id}")
async def process_section(session_id: str, request: SectionRequest,current_user: dict = Depends(get_current_user)):
    section = request.section
    proposal_id = request.proposal_id
    session_data_str = redis_client.get(session_id)

    if not session_data_str:
        raise HTTPException(status_code=400, detail="Base data not found. Please store it first.")

    with engine.connect() as connection:
        result = connection.execute(
            text("SELECT is_accepted FROM proposals WHERE id = :proposal_id AND user_id = :user_id"),
            {"proposal_id": request.proposal_id, "user_id": current_user["user_id"]}
        )
        row = result.fetchone()
        if row and row[0]:  # is_accepted is True
            raise HTTPException(status_code=403, detail="This proposal is finalized and cannot be modified.")

    session_data = json.loads(session_data_str)
    form_data = session_data["form_data"]
    project_description = session_data["project_description"]

    section_data = next((s for s in proposal_data["sections"] if s["section_name"] == section), {})
    if not section_data:
        raise HTTPException(status_code=400, detail=f"Invalid section name: {section}")

    instructions = section_data.get("instructions", "")
    word_limit = section_data.get("word_limit", 350)

    proposal_crew = ProposalCrew()
    crew_instance = proposal_crew.generate_proposal_crew()
    print("Loaded Knowledge Sources:", crew_instance.knowledge_sources[0].chunks[:1])
    pprint.pprint(crew_instance.knowledge_sources[0].chunks[:1])

    section_input = {
        "section": section,
        "form_data": form_data,
        "project_description": project_description,
        "instructions": instructions,
        "word_limit": word_limit
    }

    result = crew_instance.kickoff(inputs=section_input)
    raw_output = result.raw.replace("`", "")
    raw_output = re.sub(r'[\x00-\x1F\x7F]', '', raw_output)  
    parsed = json.loads(raw_output)


    try:
        parsed = json.loads(raw_output)
        generated_text = parsed.get("generated_content", "").strip()
        evaluation_status = parsed.get("evaluation_status", "")
        feedback = parsed.get("feedback", "")
    except Exception as e:
        return JSONResponse(
            status_code=500,
            content={"error": "Invalid JSON response from crew", "details": str(e)}
        )

    if not result:
        raise HTTPException(status_code=500, detail=f"Failed to generate content for {section}")

    is_flagged = evaluation_status.lower() == "flagged"
    evaluator_feedback = feedback.strip()

    if is_flagged and evaluator_feedback:
        generated_text = regenerate_section_logic(
            session_id=session_id,
            section=section,
            concise_input=evaluator_feedback,
            proposal_id=proposal_id
        )
        message = f"Initial content flagged. Regenerated using evaluator feedback for {section}"
    else:
        message = f"Content generated for {section}"
    session_data.setdefault("generated_sections", {})[section] = generated_text
    redis_client.set(session_id, json.dumps(session_data))
    
    #  Save section to PostgreSQL proposal (persistent)
    try:
            with engine.begin() as connection:
                result = connection.execute(
                    text("SELECT generated_sections FROM proposals WHERE id = :proposal_id AND user_id = :user_id"),
                    {
                        "proposal_id": proposal_id,
                        "user_id": current_user["user_id"]
                    }
                )
                draft = result.fetchone()

                
                if draft:
                    # Handle the case where draft[0] could be a string or already a dict
                    draft_value = draft[0]
                    
                    # Ensure it's a dictionary
                    if isinstance(draft_value, str):
                        try:
                            updated_sections = json.loads(draft_value)
                        except Exception as e:
                            print(f"[PARSE ERROR] Couldn't parse string as JSON: {e}")
                            updated_sections = {}
                    elif isinstance(draft_value, dict):
                        updated_sections = draft_value
                    else:
                        updated_sections = {}

                    # Add the new section
                    updated_sections[section] = generated_text

                    connection.execute(
                        text("""
                            UPDATE proposals
                            SET generated_sections = :sections, updated_at = NOW()
                            WHERE id = :proposal_id AND user_id = :user_id
                        """),
                        {
                            "sections": json.dumps(updated_sections),  # Always serialize to JSON string
                            "proposal_id": proposal_id,
                            "user_id": current_user["user_id"]
                        }
                    )
    except Exception as e:
            print(f"[DB UPDATE ERROR - generated_sections] {e}")



    return {
        "message": message,
        "generated_text": generated_text
    } 


@app.post("/api/regenerate_section/{session_id}")
async def regenerate_section(session_id: str,request: RegenerateRequest, current_user: dict = Depends(get_current_user)):
    """Handles regeneration of a section using a concise input for refinement."""

    section = request.section
    concise_input = request.concise_input
    proposal_id = request.proposal_id

    # ✅ Fetch session data from Redis
    session_data = redis_client.get(session_id)

    if not session_data:
        raise HTTPException(status_code=400, detail="Base data not found. Please store it first.")

    with engine.connect() as connection:
        result = connection.execute(
            text("SELECT is_accepted FROM proposals WHERE id = :proposal_id AND user_id = :user_id"),
            {"proposal_id": request.proposal_id, "user_id": current_user["user_id"]}
        )
        row = result.fetchone()
        if row and row[0]:  # is_accepted is True
            raise HTTPException(status_code=403, detail="This proposal is finalized and cannot be modified.")

    # Convert stored JSON string to dictionary
    session_data = json.loads(session_data)

    # Fetch base context
    form_data = session_data["form_data"]
    project_description = session_data["project_description"]

    # Fetch existing instructions & word limit
    section_data = next((s for s in proposal_data["sections"] if s["section_name"] == section), {})

    if not section_data:
        raise HTTPException(status_code=400, detail=f"Invalid section name: {section}")

    instructions = section_data.get("instructions", "")
    word_limit = section_data.get("word_limit", 350)

    proposal_crew = ProposalCrew()
    crew_instance = proposal_crew.regenerate_proposal_crew()

    section_input = {
        "section": section,
        "form_data": form_data,
        "project_description": project_description,
        "instructions": instructions,
        "word_limit": word_limit,
        "concise_input": concise_input
    }

    result = crew_instance.kickoff(inputs=section_input)
    raw_output = result.raw.replace("`", "")
    raw_output = re.sub(r'[\x00-\x1F\x7F]', '', raw_output)  
    parsed = json.loads(raw_output)

    try:
        parsed = json.loads(raw_output)
        generated_text = parsed.get("generated_content", "").strip()
        evaluation_status = parsed.get("evaluation_status", "")
        feedback = parsed.get("feedback", "")
    except Exception as e:
        return JSONResponse(
            status_code=500,
            content={"error": "Invalid JSON response from crew", "details": str(e)}
        )
    

    if not result:
        raise HTTPException(status_code=500, detail=f"Failed to regenerate content for {section}")

    if "generated_sections" not in session_data:
        session_data["generated_sections"] = {}  # Initialize if missing
    
    session_data["generated_sections"][section] = generated_text
    redis_client.set(session_id, json.dumps(session_data))

    user_id = current_user.get("user_id")
    if not user_id:
        raise HTTPException(status_code=400, detail="User ID missing for database update")

    try:
        with engine.begin() as connection:
            result = connection.execute(
                text("SELECT generated_sections FROM proposals WHERE id = :proposal_id AND user_id = :user_id"),
                {"proposal_id": proposal_id, "user_id": user_id}
            )
            draft = result.fetchone()

            if draft:
                draft_value = draft[0]
                
                # 💡 Ensure it's a dictionary
                if isinstance(draft_value, str):
                    try:
                        updated_sections = json.loads(draft_value)
                    except Exception as e:
                        print(f"[PARSE ERROR] Couldn't parse string as JSON: {e}")
                        updated_sections = {}
                elif isinstance(draft_value, dict):
                    updated_sections = draft_value
                else:
                    updated_sections = {}

                updated_sections[section] = generated_text

                connection.execute(
                    text("""
                        UPDATE proposals
                        SET generated_sections = :sections, updated_at = NOW()
                        WHERE id = :proposal_id AND user_id = :user_id
                    """),
                    {
                        "sections": json.dumps(updated_sections),  # ✅ Always a dict here
                        "proposal_id": proposal_id,
                        "user_id": current_user["user_id"]
                    }
                )
                
    
    except Exception as e:
        print(f"[DB UPDATE ERROR - generated_sections] {e}")

    return {
        "message": f"Content regenerated for {section}",
        "generated_text": generated_text
    }


# @app.post("/api/signup")
# async def signup(request: Request):
#     data = await request.json()
#     name = data.get('username')
#     email = data.get('email')
#     password = data.get('password')

#     if not name or not email or not password:
#         return JSONResponse(status_code=400, content={"error": "Username, Email, and Password are required."})

#     hashed_password = generate_password_hash(password)

#     try:
#         with engine.begin() as connection:
#             # Check if user exists by email
#             result = connection.execute(
#                 text("SELECT id FROM users WHERE email = :email"),
#                 {'email': email}
#             )
#             existing_user = result.fetchone()

#             if existing_user:
#                 return JSONResponse(status_code=400, content={"error": "User with this email already exists. Please log in."})

#             # Insert user into the table
#             connection.execute(
#             text("INSERT INTO users (id, email, name, password) VALUES (:id, :email, :name, :password)"),
#             {
#                 'id': str(uuid.uuid4()),
#                 'email': email,
#                 'name': name,
#                 'password': hashed_password  
#             }
#         )
#         connection.commit()

#         return JSONResponse(status_code=201, content={"message": "Signup successful! Please log in."})

#     except Exception as e:
#         print(f"[SIGNUP ERROR] {e}")
#         return JSONResponse(status_code=500, content={"error": "Signup failed. Please try again later."})

#new signup api with security questions for forgot password functionality
@app.post("/api/signup")
async def signup(request: Request):
    data = await request.json()
    name = data.get('username')
    email = data.get('email')
    password = data.get('password')
    security_question = data.get('security_question')
    security_answer = data.get('security_answer')

    # Validate all required fields
    if not name or not email or not password or not security_question or not security_answer:
        return JSONResponse(
            status_code=400,
            content={"error": "All fields are required."}
        )

    # Hash password and security answer
    hashed_password = generate_password_hash(password)
    hashed_questions = {
        security_question: generate_password_hash(security_answer.strip().lower())
    }

    try:
        with engine.begin() as connection:
            # Check if user already exists
            result = connection.execute(
                text("SELECT id FROM users WHERE email = :email"),
                {'email': email}
            )
            if result.fetchone():
                return JSONResponse(
                    status_code=400,
                    content={"error": "User with this email already exists. Please log in."}
                )

            # Insert user into the database
            connection.execute(
                text("""
                    INSERT INTO users (id, email, name, password, security_questions)
                    VALUES (:id, :email, :name, :password, :security_questions)
                """),
                {
                    'id': str(uuid.uuid4()),
                    'email': email,
                    'name': name,
                    'password': hashed_password,
                    'security_questions': json.dumps(hashed_questions)
                }
            )

        return JSONResponse(status_code=201, content={"message": "Signup successful! Please log in."})

    except Exception as e:
        print(f"[SIGNUP ERROR] {e}")
        return JSONResponse(status_code=500, content={"error": "Signup failed. Please try again later."})


@app.post("/api/login")
async def login(request: Request):
    try:
        data = await request.json()
        email = data.get('email')
        password = data.get('password')

        if not email or not password:
            return JSONResponse(status_code=400, content={"error": "Email and password are required."})

        with engine.connect() as connection:
            result = connection.execute(text("SELECT id, email, name, password FROM users WHERE email = :email"), {'email': email})
            print("📧 Incoming email for login:", email)

            user = result.fetchone()
            print("🧑‍💻 DB query result:", user)


        if not user:
            return JSONResponse(status_code=404, content={"error": "User does not exist!"})

        user_id = user[0]
        stored_password = user[3]  # password is 3rd column

        if not check_password_hash(stored_password, password):
            return JSONResponse(status_code=401, content={"error": "Invalid password!"})

        # ✅ Check if user already has an active session
        existing_session_token = redis_client.get(f"user_session:{user_id}")
        if existing_session_token:
            return JSONResponse(
                status_code=403,
                content={"error": "You are already logged in from another session."}
            )
        
        # ✅ Create JWT token
        token = jwt.encode(
            {
                "email": email,
                "exp": datetime.utcnow() + timedelta(minutes=30)
            },
            SECRET_KEY,
            algorithm="HS256"
        )
        
        # ✅ Store token in Redis for this user
        redis_client.setex(
            f"user_session:{user_id}",
            1800,  # 30 minutes
            token
        )

        # ✅ Set token in HttpOnly cookie
        response = JSONResponse(content={"message": "Login successful!"})
        response.set_cookie(
            key="auth_token",
            value=token,
            httponly=True,
            samesite="Lax",
            max_age=1800  # 30 minutes
        )

        return response

    except Exception as e:
        print(f"[LOGIN ERROR] {e}")
        return JSONResponse(status_code=500, content={"error": "Login failed. Please try again later."})


@app.get("/api/profile")
async def profile(current_user: dict = Depends(get_current_user)):
    return {
        "message": "Profile fetched successfully",
        "user": {
            "name": current_user["name"],
            "email": current_user["email"]
        }
    }

@app.post("/api/logout")
async def logout(current_user: dict = Depends(get_current_user)):
    """Logs out the user by removing the auth_token cookie and Redis session."""
    user_id = current_user["user_id"]

    try:
        # ✅ Delete active session token from Redis
        redis_client.delete(f"user_session:{user_id}")
        print(f"[LOGOUT] Removed session for user_id: {user_id}")

    except Exception as e:
        print(f"[LOGOUT ERROR] Failed to remove session for user_id {user_id}: {e}")

    # ✅ Remove cookie
    response = JSONResponse(content={"message": "Logout successful!"})
    response.delete_cookie(key="auth_token")

    return response

@app.post("/api/get-security-question")
async def get_security_question(request: Request):
    try:
        data = await request.json()
        email = data.get("email")

        if not email:
            return JSONResponse(
                status_code=400,
                content={"error": "Email is required."}
            )

        with engine.connect() as connection:
            result = connection.execute(
                text("SELECT security_questions FROM users WHERE email = :email"),
                {"email": email}
            )
            user = result.fetchone()

        if not user:
            return JSONResponse(status_code=404, content={"error": "User not found."})

        stored_questions_json = user[0]
        if not stored_questions_json:
            return JSONResponse(status_code=400, content={"error": "Security question not set."})

        # Handle possible JSON format
        if isinstance(stored_questions_json, str):
            stored_questions = json.loads(stored_questions_json)
        elif isinstance(stored_questions_json, dict):
            stored_questions = stored_questions_json
        else:
            return JSONResponse(status_code=500, content={"error": "Invalid format of stored security questions."})

        # Extract the only question
        if len(stored_questions) != 1:
            return JSONResponse(status_code=500, content={"error": "Invalid number of stored security questions."})

        question = list(stored_questions.keys())[0]
        return JSONResponse(status_code=200, content={"question": question})

    except Exception as e:
        print(f"[GET SECURITY QUESTION ERROR] {e}")
        return JSONResponse(status_code=500, content={"error": "Something went wrong. Please try again later."})


@app.post("/api/verify-security-answer")
async def verify_security_answer(request: Request):
    try:
        data = await request.json()
        email = data.get("email")
        security_question = data.get("security_question")
        security_answer = data.get("security_answer")

        if not email or not security_question or not security_answer:
            return JSONResponse(
                status_code=400,
                content={"error": "All fields are required."}
            )

        with engine.connect() as connection:
            result = connection.execute(
                text("SELECT security_questions FROM users WHERE email = :email"),
                {"email": email}
            )
            user = result.fetchone()

            if not user:
                return JSONResponse(status_code=404, content={"error": "User not found."})

            stored_questions_json = user[0]
            if isinstance(stored_questions_json, str):
                stored_questions = json.loads(stored_questions_json)
            elif isinstance(stored_questions_json, dict):
                stored_questions = stored_questions_json
            else:
                return JSONResponse(status_code=500, content={"error": "Invalid format of stored security questions."})

            hashed_answer = stored_questions.get(security_question)
            if not hashed_answer:
                return JSONResponse(status_code=403, content={"error": "Security question mismatch."})

            if not check_password_hash(hashed_answer, security_answer.strip().lower()):
                return JSONResponse(status_code=403, content={"error": "Incorrect security answer."})

        # Security answer verified successfully
        return JSONResponse(status_code=200, content={"message": "Security answer verified successfully."})

    except Exception as e:
        print(f"[VERIFY SECURITY ANSWER ERROR] {e}")
        return JSONResponse(status_code=500, content={"error": "Failed to verify security answer. Please try again later."})


@app.post("/api/update-password")
async def update_password(request: Request):
    try:
        data = await request.json()
        email = data.get("email")
        security_question = data.get("security_question")
        security_answer = data.get("security_answer")
        new_password = data.get("new_password")

        if not email or not security_question or not security_answer or not new_password:
            return JSONResponse(
                status_code=400,
                content={"error": "All fields are required."}
            )

        with engine.begin() as connection:
            # Verify security answer again for security
            result = connection.execute(
                text("SELECT security_questions FROM users WHERE email = :email"),
                {"email": email}
            )
            user = result.fetchone()

            if not user:
                return JSONResponse(status_code=404, content={"error": "User not found."})

            stored_questions_json = user[0]
            if isinstance(stored_questions_json, str):
                stored_questions = json.loads(stored_questions_json)
            elif isinstance(stored_questions_json, dict):
                stored_questions = stored_questions_json
            else:
                return JSONResponse(status_code=500, content={"error": "Invalid format of stored security questions."})

            hashed_answer = stored_questions.get(security_question)
            if not hashed_answer:
                return JSONResponse(status_code=403, content={"error": "Security question mismatch."})

            if not check_password_hash(hashed_answer, security_answer.strip().lower()):
                return JSONResponse(status_code=403, content={"error": "Incorrect security answer."})

            # All good — update password
            hashed_password = generate_password_hash(new_password)
            connection.execute(
                text("UPDATE users SET password = :password WHERE email = :email"),
                {"password": hashed_password, "email": email}
            )

        return JSONResponse(status_code=200, content={"message": "Password updated successfully. Please log in."})

    except Exception as e:
        print(f"[UPDATE PASSWORD ERROR] {e}")
        return JSONResponse(status_code=500, content={"error": "Failed to update password. Please try again later."})

@app.post("/api/save-draft")
async def save_draft(request: SaveDraftRequest, current_user: dict = Depends(get_current_user)):
    try:
        user_id = current_user["user_id"]
        proposal_id = request.proposal_id or str(uuid.uuid4())
        
        with engine.begin() as connection:
            result = connection.execute(
                text("SELECT id FROM proposals WHERE id = :proposal_id AND user_id = :user_id"),
                {"proposal_id": proposal_id, "user_id": user_id}
            )
            existing = result.fetchone()

            if existing:
                # Update existing proposal
                connection.execute(
                    text("""
                        UPDATE proposals 
                        SET form_data = :form_data,
                            project_description = :project_description,
                            generated_sections = :sections,
                            updated_at = NOW()
                        WHERE id = :proposal_id AND user_id = :user_id
                    """),
                    {
                        "form_data": json.dumps(request.form_data),
                        "project_description": request.project_description,  # ✅ Match bind key
                        "sections": json.dumps(request.generated_sections),
                        "proposal_id": proposal_id,
                        "user_id": user_id
                    }
                )

                message = "Draft updated successfully"
            else:
                # Insert new proposal
                connection.execute(
                    text("""
                        INSERT INTO proposals 
                        (id, user_id, form_data, project_description, generated_sections)
                        VALUES (:id, :user_id, :form_data, :project_description, :sections)
                    """),
                    {
                        "id": proposal_id,
                        "user_id": user_id,
                        "form_data": json.dumps(request.form_data),
                        "project_description": request.project_description,  # ✅ Correct key
                        "sections": json.dumps(request.generated_sections)
                    }

                )
                message = "Draft created successfully"
        if request.session_id:
            try:
                redis_data_raw = redis_client.get(request.session_id)
                if redis_data_raw:
                    if isinstance(redis_data_raw, bytes):
                        redis_data_raw = redis_data_raw.decode("utf-8")
                    redis_data = json.loads(redis_data_raw)
                    redis_data["proposal_id"] = proposal_id
                    redis_client.setex(request.session_id, 3600, json.dumps(redis_data))
                    print(f"[SAVE-DRAFT] Updated Redis session {request.session_id} with proposal_id: {proposal_id}")
            except Exception as e:
                print(f"[SAVE-DRAFT REDIS UPDATE ERROR] {e}")

        return {
            "message": message,
            "proposal_id": proposal_id
        }

    except Exception as e:
        print(f"[SAVE DRAFT ERROR] {e}")
        raise HTTPException(status_code=500, detail="Failed to save draft")



# @app.get("/api/list-drafts")
# async def list_drafts(current_user: dict = Depends(get_current_user)):
#     user_id = current_user["user_id"]

#     try:
#         with engine.connect() as connection:
#             # ✅ Include form_data in the query itself
#             result = connection.execute(
#                 text("""
#                     SELECT id, form_data, generated_sections, created_at, updated_at, is_accepted
#                     FROM proposals
#                     WHERE user_id = :user_id
#                     ORDER BY updated_at DESC
#                 """),
#                 {"user_id": user_id}
#             )
#             drafts = result.fetchall()

#         if not drafts:
#             return {
#                 "message": "No drafts found.",
#                 "drafts": []
#             }

#         draft_list = []

#         for row in drafts:
#             proposal_id = row[0]
#             form_data_raw = row[1]
#             generated_sections_raw = row[2]
#             created_at = row[3]
#             updated_at = row[4]
#             is_accepted = row[5] if len(row) > 5 else False


#             # ✅ Parse form_data
#             try:
#                 if isinstance(form_data_raw, str):
#                     form_data = json.loads(form_data_raw)
#                 elif isinstance(form_data_raw, dict):
#                     form_data = form_data_raw
#                 else:
#                     form_data = {}
#             except Exception as e:
#                 print(f"[FORM_DATA PARSE ERROR for {proposal_id}] {e}")
#                 form_data = {}

#             project_title = form_data.get("Project title", "Untitled Proposal")

#             # ✅ Parse generated_sections
#             try:
#                 if isinstance(generated_sections_raw, str):
#                     generated_sections = json.loads(generated_sections_raw)
#                 elif isinstance(generated_sections_raw, dict):
#                     generated_sections = generated_sections_raw
#                 else:
#                     generated_sections = {}
#             except Exception as e:
#                 print(f"[PARSE ERROR] Could not parse generated_sections for {proposal_id}: {e}")
#                 generated_sections = {}

#             summary = generated_sections.get("Summary")

#             draft_list.append({
#                 "proposal_id": proposal_id,
#                 "project_title": project_title,
#                 "summary": summary,
#                 "is_accepted": is_accepted,
#                 "created_at": created_at.isoformat() if created_at else None,
#                 "updated_at": updated_at.isoformat() if updated_at else None
#             })

#         return {
#             "message": "Drafts fetched successfully.",
#             "drafts": draft_list
#         }

#     except Exception as e:
#         print(f"[LIST DRAFTS ERROR] {e}")
#         raise HTTPException(status_code=500, detail="Failed to fetch drafts")


# @app.get("/api/load-draft/{proposal_id}")
# async def load_draft(proposal_id: str, current_user: dict = Depends(get_current_user)):
#     user_id = current_user["user_id"]

#     try:
#         # ✅ Check if a Redis session already exists for this proposal
#         redis_keys = redis_client.keys("*")
#         for key in redis_keys:
#             data = redis_client.get(key)
#             if not data:
#                 continue

#             try:
#                 if isinstance(data, bytes):
#                     data = data.decode("utf-8")

#                 if data.strip():  # Only try parsing if data is non-empty
#                     parsed = json.loads(data)
#                 else:
#                     continue
#             except Exception as e:
#                 print(f"[REDIS JSON PARSE ERROR for key={key}] {e}")
#                 continue  # Skip this corrupted session

#             if parsed.get("user_id") == user_id and parsed.get("proposal_id") == proposal_id:
#                     print(f"[LOAD DRAFT] Reusing existing Redis session: {key}")
#                     return {
#                         "form_data": parsed["form_data"],
#                         "project_description": parsed["project_description"],
#                         "generated_sections": {section: parsed.get("generated_sections", {}).get(section) for section in SECTIONS},
#                         "session_id": key,
#                         "proposal_id": proposal_id,
#                         "is_accepted": parsed.get("is_accepted", False) 
#                     }
                
#         with engine.connect() as connection:
#             result = connection.execute(
#                 text("""
#                     SELECT form_data, project_description, generated_sections, is_accepted
#                     FROM proposals
#                     WHERE id = :proposal_id AND user_id = :user_id
#                 """),
#                 {"proposal_id": proposal_id, "user_id": user_id}
#             )
#             draft = result.fetchone()

#         if not draft:
#             raise HTTPException(status_code=404, detail="Draft not found for this proposal_id")

#         form_data_raw, project_description, generated_sections_raw, is_accepted = draft

#         # Parse loaded data
#         form_data = json.loads(draft[0]) if isinstance(draft[0], str) else draft[0]
#         project_description = draft[1]
#         loaded_sections = json.loads(draft[2]) if draft[2] and isinstance(draft[2], str) else draft[2] or {}
#         ordered_sections = {section: loaded_sections.get(section) for section in SECTIONS}
#         # form_data, project_description, generated_sections_raw, is_accepted = draft
        
#         # ✅ Create a new session_id
#         session_id = str(uuid.uuid4())

#         # ✅ Push form_data and project_description to Redis under new session_id
#         redis_client.setex(
#             session_id,
#             3600,  # 1 hour expiry
#             json.dumps({
#                 "form_data": form_data,
#                 "project_description": project_description,
#                 "generated_sections": ordered_sections,
#                 "project_title": form_data.get("project_title", "Untitled Proposal"),
#                 "user_id": user_id,
#                 "proposal_id": proposal_id,
#                 "is_accepted": is_accepted,
#             })
#         )

#         print(f"[LOAD DRAFT] Created new Redis session: {session_id}")

#         # ✅ Return data including new session_id
#         return {
#             "form_data": form_data,
#             "project_description": project_description,
#             "generated_sections": ordered_sections,
#             "session_id": session_id,  
#             "proposal_id": proposal_id,
#             "is_accepted": is_accepted
#         }

#     except Exception as e:
#         print(f"[LOAD DRAFT ERROR] {e}")
#         raise HTTPException(status_code=404, detail="Draft not found")

# @app.get("/api/load-draft/{proposal_id}")
# async def load_draft(proposal_id: str, current_user: dict = Depends(get_current_user)):
#     user_id = current_user["user_id"]

#     try:
#         with engine.connect() as connection:
#             # ✅ Always load latest data directly from DB (not Redis)
#             result = connection.execute(
#                 text("""
#                     SELECT form_data, project_description, generated_sections, is_accepted, created_at, updated_at
#                     FROM proposals
#                     WHERE id = :proposal_id AND user_id = :user_id
#                 """),
#                 {"proposal_id": proposal_id, "user_id": user_id}
#             )
#             draft = result.fetchone()

#         if not draft:
#             raise HTTPException(status_code=404, detail="Draft not found for this proposal_id")

#         form_data_raw, project_description, generated_sections_raw, is_accepted, created_at, updated_at = draft

#         # ✅ Parse form_data
#         try:
#             if isinstance(form_data_raw, str):
#                 form_data = json.loads(form_data_raw)
#             elif isinstance(form_data_raw, dict):
#                 form_data = form_data_raw
#             else:
#                 form_data = {}
#         except Exception as e:
#             print(f"[FORM_DATA PARSE ERROR] {e}")
#             form_data = {}

#         # ✅ Parse generated_sections
#         try:
#             if isinstance(generated_sections_raw, str):
#                 generated_sections = json.loads(generated_sections_raw)
#             elif isinstance(generated_sections_raw, dict):
#                 generated_sections = generated_sections_raw
#             else:
#                 generated_sections = {}
#         except Exception as e:
#             print(f"[GENERATED SECTIONS PARSE ERROR] {e}")
#             generated_sections = {}

#         # ✅ Ensure all expected sections are present (as per SECTIONS list)
#         ordered_sections = {section: generated_sections.get(section) for section in SECTIONS}

#         # ✅ Create new session and store in Redis (optional, for continuity)
#         session_id = str(uuid.uuid4())
#         redis_client.setex(
#             session_id,
#             3600,  # 1 hour expiry
#             json.dumps({
#                 "form_data": form_data,
#                 "project_description": project_description,
#                 "generated_sections": ordered_sections,
#                 "project_title": form_data.get("Project title", "Untitled Proposal"),
#                 "user_id": user_id,
#                 "proposal_id": proposal_id,
#                 "is_accepted": is_accepted,
#             })
#         )
#         print(f"[LOAD DRAFT] Redis session created: {session_id}")

#         # ✅ Return final payload (no dependency on Redis for loading)
#         return {
#             "form_data": form_data,
#             "project_description": project_description,
#             "generated_sections": ordered_sections,
#             "session_id": session_id,
#             "proposal_id": proposal_id,
#             "is_accepted": is_accepted,
#             "created_at": created_at.isoformat() if created_at else None,
#             "updated_at": updated_at.isoformat() if updated_at else None
#         }

#     except Exception as e:
#         print(f"[LOAD DRAFT ERROR] {e}")
#         raise HTTPException(status_code=500, detail="Failed to load draft")

@app.get("/api/list-drafts")
async def list_drafts(current_user: dict = Depends(get_current_user)):
    user_id = current_user["user_id"]

    draft_list = []

    #  Load sample templates first
    try:
        with open("config/templates/sample_templates.json", "r", encoding="utf-8") as f:
            sample_templates = json.load(f)
    except Exception as e:
        print(f"[TEMPLATE LOAD ERROR] Failed to load sample templates: {e}")
        sample_templates = []

    # Normalize sample templates to match frontend format
    for sample in sample_templates:
        form_data = sample.get("form_data", {})
        sections = sample.get("generated_sections", {})

        sample["project_title"] = form_data.get("Project title", "Untitled Proposal")
        sample["summary"] = sections.get("Summary", "")
        sample["is_accepted"] = sample.get("is_accepted", True)
        sample["created_at"] = sample.get("created_at", datetime.utcnow().isoformat())
        sample["updated_at"] = sample.get("updated_at", datetime.utcnow().isoformat())
        sample["is_sample"] = True  # ensure flag is explicitly set

    draft_list.extend(sample_templates)

    try:
        with engine.connect() as connection:
            result = connection.execute(
                text("""
                    SELECT id, form_data, generated_sections, created_at, updated_at, is_accepted
                    FROM proposals
                    WHERE user_id = :user_id
                    ORDER BY updated_at DESC
                """),
                {"user_id": user_id}
            )
            drafts = result.fetchall()

        for row in drafts:
            proposal_id = row[0]
            form_data_raw = row[1]
            generated_sections_raw = row[2]
            created_at = row[3]
            updated_at = row[4]
            is_accepted = row[5] if len(row) > 5 else False

            # Parse form_data
            try:
                form_data = json.loads(form_data_raw) if isinstance(form_data_raw, str) else form_data_raw or {}
            except Exception as e:
                print(f"[FORM_DATA PARSE ERROR for {proposal_id}] {e}")
                form_data = {}

            # Parse generated_sections
            try:
                generated_sections = json.loads(generated_sections_raw) if isinstance(generated_sections_raw, str) else generated_sections_raw or {}
            except Exception as e:
                print(f"[PARSE ERROR] Could not parse generated_sections for {proposal_id}: {e}")
                generated_sections = {}

            project_title = form_data.get("Project title", "Untitled Proposal")
            summary = generated_sections.get("Summary", "")

            draft_list.append({
                "proposal_id": proposal_id,
                "form_data": form_data,
                "generated_sections": generated_sections,
                "project_title": project_title,
                "summary": summary,
                "is_accepted": is_accepted,
                "created_at": created_at.isoformat() if created_at else None,
                "updated_at": updated_at.isoformat() if updated_at else None,
                "is_sample": False
            })

        return {
            "message": "Drafts fetched successfully.",
            "drafts": draft_list
        }

    except Exception as e:
        print(f"[LIST DRAFTS ERROR] {e}")
        raise HTTPException(status_code=500, detail="Failed to fetch drafts")

@app.get("/api/load-draft/{proposal_id}")
async def load_draft(proposal_id: str, current_user: dict = Depends(get_current_user)):
    user_id = current_user["user_id"]

    # Check if it's a sample draft
    if proposal_id.startswith("sample-"):
        try:
            with open("config/templates/sample_templates.json", "r", encoding="utf-8") as f:
                sample_templates = json.load(f)

            sample = next((item for item in sample_templates if item.get("proposal_id") == proposal_id), None)

            if not sample:
                raise HTTPException(status_code=404, detail="Sample draft not found.")

            form_data = sample.get("form_data", {})
            generated_sections = sample.get("generated_sections", {})
            project_description = generated_sections.get("Project Description", "")

            session_id = str(uuid.uuid4())
            redis_client.setex(
                session_id,
                3600,
                json.dumps({
                    "form_data": form_data,
                    "project_description": project_description,
                    "generated_sections": generated_sections,
                    "project_title": form_data.get("Project title", "Untitled Proposal"),
                    "user_id": user_id,
                    "proposal_id": proposal_id,
                    "is_accepted": sample.get("is_accepted", True),
                    "is_sample": True
                })
            )

            return {
                "form_data": form_data,
                "project_description": project_description,
                "generated_sections": generated_sections,
                "session_id": session_id,
                "proposal_id": proposal_id,
                "is_accepted": sample.get("is_accepted", True),
                "is_sample": True,
                "created_at": sample.get("created_at"),
                "updated_at": sample.get("updated_at")
            }

        except Exception as e:
            print(f"[SAMPLE LOAD ERROR] {e}")
            raise HTTPException(status_code=500, detail="Failed to load sample draft")

    # Fallback: regular draft from DB
    try:
        with engine.connect() as connection:
            result = connection.execute(
                text("""
                    SELECT form_data, project_description, generated_sections, is_accepted, created_at, updated_at
                    FROM proposals
                    WHERE id = :proposal_id AND user_id = :user_id
                """),
                {"proposal_id": proposal_id, "user_id": user_id}
            )
            draft = result.fetchone()

        if not draft:
            raise HTTPException(status_code=404, detail="Draft not found for this proposal_id")

        form_data_raw, project_description, generated_sections_raw, is_accepted, created_at, updated_at = draft

        try:
            form_data = json.loads(form_data_raw) if isinstance(form_data_raw, str) else form_data_raw or {}
        except Exception as e:
            print(f"[FORM_DATA PARSE ERROR] {e}")
            form_data = {}

        try:
            generated_sections = json.loads(generated_sections_raw) if isinstance(generated_sections_raw, str) else generated_sections_raw or {}
        except Exception as e:
            print(f"[GENERATED SECTIONS PARSE ERROR] {e}")
            generated_sections = {}

        ordered_sections = {section: generated_sections.get(section) for section in SECTIONS}

        session_id = str(uuid.uuid4())
        redis_client.setex(
            session_id,
            3600,
            json.dumps({
                "form_data": form_data,
                "project_description": project_description,
                "generated_sections": ordered_sections,
                "project_title": form_data.get("Project title", "Untitled Proposal"),
                "user_id": user_id,
                "proposal_id": proposal_id,
                "is_accepted": is_accepted,
                "is_sample": False
            })
        )

        return {
            "form_data": form_data,
            "project_description": project_description,
            "generated_sections": ordered_sections,
            "session_id": session_id,
            "proposal_id": proposal_id,
            "is_accepted": is_accepted,
            "is_sample": False,
            "created_at": created_at.isoformat() if created_at else None,
            "updated_at": updated_at.isoformat() if updated_at else None
        }

    except Exception as e:
        raise HTTPException(status_code=500, detail="Failed to load draft: {str(e)}")
 

def add_markdown_paragraph(doc, text):
    paragraph = doc.add_paragraph()
    # Split on things wrapped with **
    parts = re.split(r'(\*\*.*?\*\*)', text)
    
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            paragraph.add_run(part[2:-2]).bold = True
        else:
            paragraph.add_run(part)
    
    # Optional formatting
    paragraph.paragraph_format.space_after = Pt(12)
    paragraph.paragraph_format.line_spacing = 1.5

def convert_markdown_bold(text):
    """
    Safely converts **bold** markdown to <b>bold</b> for ReportLab.
    """
    return re.sub(r"\*\*(.+?)\*\*", r"<b>\1</b>", text)

def create_pdf_from_sections(output_path, form_data, ordered_sections):
    
    doc = SimpleDocTemplate(output_path, pagesize=A4, rightMargin=72, leftMargin=72, topMargin=72, bottomMargin=72)
    styles = getSampleStyleSheet()

    # Custom paragraph style for justified text
    normal_style = ParagraphStyle(
        name='Body',
        parent=styles['Normal'],
        alignment=TA_JUSTIFY,
        leading=14
    )

    story = []

    # Title
    story.append(Paragraph("Project Proposal", styles['Title']))
    story.append(Spacer(1, 12))

    # Build table data from form_data
    table_data = [["Field", "Value"]]
    for key, value in form_data.items():
        table_data.append([key, value])

    table = Table(table_data, colWidths=[2.5 * inch, 3.5 * inch])
    table.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, 0), colors.lightgrey),
        ('TEXTCOLOR', (0, 0), (-1, 0), colors.black),
        ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
        ('FONTSIZE', (0, 0), (-1, -1), 10),
        ('BOTTOMPADDING', (0, 0), (-1, 0), 8),
        ('GRID', (0, 0), (-1, -1), 0.5, colors.grey),
    ]))
    story.append(table)
    story.append(Spacer(1, 20))

    # Add each section
    for section, content in ordered_sections.items():
        if not content:
            continue  # skip empty sections

        story.append(Paragraph(section, styles['Heading2']))
        story.append(Spacer(1, 6))

        for paragraph in content.split("\n\n"):
            cleaned = convert_markdown_bold(paragraph.strip())
            story.append(Paragraph(cleaned, normal_style))
            story.append(Spacer(1, 10))

    # Build the PDF
    doc.build(story)



@app.get("/api/generate-document/{proposal_id}")
async def generate_and_download_document(
    proposal_id: str,
    format: str = "docx",  # either 'docx' or 'pdf'
    current_user: dict = Depends(get_current_user)
):
    """Generates final document (Word & PDF) and returns requested file as download"""

    user_id = current_user["user_id"]

    try:
        with engine.connect() as connection:
            result = connection.execute(
                text("""
                    SELECT form_data, project_description, generated_sections
                    FROM proposals
                    WHERE id = :proposal_id AND user_id = :user_id
                """),
                {"proposal_id": proposal_id, "user_id": user_id}
            )
            draft = result.fetchone()

        if not draft:
            raise HTTPException(status_code=404, detail="Proposal not found for this user.")

        form_data = json.loads(draft[0]) if isinstance(draft[0], str) else draft[0]
        # generated_sections = json.loads(draft[2]) if draft[2] and isinstance(draft[2], str) else {}
        try:
            if draft[2] and isinstance(draft[2], str):
                # Handle double-encoded JSON (common when inserted via json.dumps twice)
                first_pass = json.loads(draft[2])
                generated_sections = json.loads(first_pass) if isinstance(first_pass, str) else first_pass
            elif isinstance(draft[2], dict):
                generated_sections = draft[2]
            else:
                generated_sections = {}
        except Exception as e:
            print(f"[JSON PARSE ERROR - generated_sections] {e}")
            generated_sections = {}


        if len(generated_sections) != len(SECTIONS):
            missing_sections = [s for s in SECTIONS if s not in generated_sections]
            raise HTTPException(status_code=400, detail=f"Missing sections: {missing_sections}")

        ordered_sections = {section: generated_sections.get(section) for section in SECTIONS}

        # ✅ Create Word document
        doc = Document()
        doc.add_heading("Project Proposal", level=1)

        table = doc.add_table(rows=1, cols=2)
        table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Field'
        hdr_cells[1].text = 'Value'

        for key, value in form_data.items():
            row_cells = table.add_row().cells
            row_cells[0].text = key
            row_cells[1].text = value

        doc.add_paragraph("\n")

        for section, content in ordered_sections.items():
            doc.add_heading(section, level=2)
            for para in content.split("\n\n"):
                add_markdown_paragraph(doc, para.strip())

        # ✅ Save DOCX
        folder_name = "proposal-documents"
        os.makedirs(folder_name, exist_ok=True)
        unique_id = f"{datetime.now().strftime('%Y%m%d%H%M%S')}_{uuid.uuid4().hex[:6]}"
        docx_file_path = os.path.join(folder_name, f"proposal_{unique_id}.docx")
        
        doc.save(docx_file_path)
    
        pdf_file_path = None

        if format == "pdf":
            pdf_file_path = docx_file_path.replace(".docx", ".pdf")
            try:
                create_pdf_from_sections(pdf_file_path, form_data, ordered_sections)
                print(f"[PDF Created] {pdf_file_path}")
            except Exception as e:
                print(f"[PDF Generation Error] {e}")
                pdf_file_path = None

        # ✅ Return requested file
        if format == "pdf":
            if not pdf_file_path or not os.path.exists(pdf_file_path):
                raise HTTPException(status_code=500, detail="Failed to generate PDF document.")
            return FileResponse(
                path=pdf_file_path,
                filename=f"Proposal_{proposal_id}.pdf",
                media_type='application/pdf'
            )
        else:
            return FileResponse(
                path=docx_file_path,
                filename=f"Proposal_{proposal_id}.docx",
                media_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
            )

    except Exception as e:
        traceback.print_exc()  # 🔍 Print full traceback
        print(f"[GENERATE & DOWNLOAD ERROR] {type(e).__name__}: {e}")
        raise HTTPException(status_code=500, detail=f"Document generation failed due to: {str(e)}")



@app.post("/api/finalize-proposal")
async def finalize_proposal(request: FinalizeProposalRequest, current_user: dict = Depends(get_current_user)):
    proposal_id = request.proposal_id
    user_id = current_user["user_id"]

    try:
        with engine.begin() as connection:
            result = connection.execute(
                text("SELECT is_accepted FROM proposals WHERE id = :proposal_id AND user_id = :user_id"),
                {"proposal_id": proposal_id, "user_id": user_id}
            )
            proposal = result.fetchone()

            if not proposal:
                raise HTTPException(status_code=404, detail="Proposal not found.")

            if proposal[0]:  # already finalized
                return {"message": "Proposal is already finalized."}

            # ✅ Set is_accepted = TRUE
            connection.execute(
                text("""
                    UPDATE proposals
                    SET is_accepted = TRUE, updated_at = NOW()
                    WHERE id = :proposal_id AND user_id = :user_id
                """),
                {"proposal_id": proposal_id, "user_id": user_id}
            )

        return {
            "message": "Proposal finalized successfully.",
            "proposal_id": proposal_id,
            "is_accepted": True
        }

    except Exception as e:
        print(f"[FINALIZE PROPOSAL ERROR] {e}")
        raise HTTPException(status_code=500, detail="Failed to finalize proposal.")


@app.delete("/api/delete-draft/{proposal_id}")
async def delete_draft(
    proposal_id: str,
    # session_id: Optional[str] = Query(None),  # 👈 optional query param for Redis session_id
    current_user: dict = Depends(get_current_user)
):
    user_id = current_user["user_id"]

    try:
        with engine.begin() as connection:
            # Check if the proposal exists for this user
            result = connection.execute(
                text("SELECT id FROM proposals WHERE id = :proposal_id AND user_id = :user_id"),
                {"proposal_id": proposal_id, "user_id": user_id}
            )
            draft = result.fetchone()

            if not draft:
                raise HTTPException(status_code=404, detail="Draft not found.")

            # Delete the proposal from the DB
            connection.execute(
                text("DELETE FROM proposals WHERE id = :proposal_id AND user_id = :user_id"),
                {"proposal_id": proposal_id, "user_id": user_id}
            )

        # If a session_id is provided, delete the associated Redis session
        # Step 2: Delete any Redis sessions associated with this draft
        session_deleted = False
        redis_keys = redis_client.keys("*")
        for key in redis_keys:
            data = redis_client.get(key)
            if not data:
                continue

            try:
                parsed = json.loads(data)
                if parsed.get("user_id") == user_id and parsed.get("proposal_id") == proposal_id:
                    redis_client.delete(key)
                    session_deleted = True
                    print(f"[DELETE] Removed Redis session for session_id: {key}")
            except Exception as e:
                print(f"[REDIS CLEANUP ERROR - key={key}] {e}")

        return {
            "message": f"Draft '{proposal_id}' deleted successfully.",
            "session_deleted": session_deleted
        }

    except Exception as e:
        print(f"[DELETE DRAFT ERROR] {e}")
        raise HTTPException(status_code=500, detail="Failed to delete draft")

@app.get("/api/health_check")
def health_check():
    return {"status": "API is running"}

def delete_old_proposals():
    try:
        # ⏱ For testing: delete proposals older than 5 minutes
        threshold = datetime.now(timezone.utc) - timedelta(days=90)

        with engine.begin() as connection:
            result = connection.execute(
                text("DELETE FROM proposals WHERE created_at < :threshold"),
                {"threshold": threshold}
            )
            print(f"[CLEANUP] Deleted old proposals older than {threshold}")
    except Exception as e:
        print(f"[CLEANUP ERROR] {e}")

# 🗓 Start the background scheduler
scheduler = BackgroundScheduler()
scheduler.add_job(delete_old_proposals, 'interval', minutes=1)  # Runs every 1 minute
scheduler.start()


if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="0.0.0.0", port=8502)
